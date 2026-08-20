"""Extração DOCX."""

from __future__ import annotations

import logging
from pathlib import Path

from app.to_xlsx.extractors import ExtractionResult

logger = logging.getLogger(__name__)


def extract_docx(path: str) -> ExtractionResult:
    from docx import Document

    doc = Document(path)
    parts: list[str] = []

    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    logger.info("DOCX: %s blocos de %s", len(parts), Path(path).name)
    return ExtractionResult(
        source_format="docx",
        text=text,
        structured=False,
    )
